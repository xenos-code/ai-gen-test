# app_functions.py

import pandas as pd
import streamlit as st
import time
import re
import os
import zipfile
import openai
from prompts import prompts
from docx import Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
from docx.opc.constants import RELATIONSHIP_TYPE
from html.parser import HTMLParser

def render_expanders(expanders):
    for key, value in expanders.items():
        with st.expander(value["title"]):
            st.markdown(value["content"])

def create_url_path(keyword):
    url_path = keyword.lower()
    url_path = re.sub(r"[^a-z0-9\s]+", "", url_path)  # Remove non-alphanumeric and non-space characters
    url_path = url_path.replace(" ", "-")  # Replace spaces with hyphens

    # Remove trailing hyphen if present
    if url_path[-1] == "-":
        url_path = url_path[:-1]

    return f"/{url_path}"

def create_full_path(domain, url_path):
    return f"https://{domain}{url_path}"

def generate_content(api_key, prompt, sections, model, temperature, presence_penalty, frequency_penalty, max_tokens):
    openai.api_key = api_key

    system_message = prompts["system_message"]
    
    print(f"Generated prompt:\n{prompt}\n")

    messages = [
        {"role": "system", "content": system_message},
        {"role": "user", "content": prompt}
    ]
    
    
    print(f"API Call Parameters:")
    print(f"API Key: {api_key}")
    print(f"GPT Model: {model}")
    print(f"Prompt: {prompt}")
    print(f"Sections: {sections}")
    print(f"Temperature: {temperature}")
    print(f"Presence Penalty: {presence_penalty}")
    print(f"Frequency Penalty: {frequency_penalty}")
    print(f"Max Tokens: {max_tokens}")

    completion = openai.ChatCompletion.create(
        model=model,
        max_tokens=max_tokens,
        temperature=temperature,
        presence_penalty=presence_penalty,
        frequency_penalty=frequency_penalty,
        messages=messages,
    )

    response = completion.choices[0].message.content.strip()
    # print(f"Generated response:\n{response}\n")
    return response

def generate_related_links(df, current_topic):
    current_category = df.loc[df['topic'] == current_topic, 'category'].values[0]
    current_full_path = df.loc[df['topic'] == current_topic, 'full path'].values[0]
    related_links = df[df['category'] == current_category][['topic', 'full path']]
    
    # Filter out the self-link by comparing the full paths
    related_links = related_links[related_links['full path'] != current_full_path]

    return related_links.to_dict('records')

def generate_article(api_key, topic, sections, related_links, model, temperature, presence_penalty, frequency_penalty, max_tokens, definition_only=False):
    if definition_only:
        prompt = prompts["definition_prompt"].format(topic)
    else:
        if related_links:
            related_links_prompt = prompts["related_links_prompt"].format(
                ", ".join([f"{rl['topic']} ({rl['full path']})" for rl in related_links])
            )
        else:
            related_links_prompt = ""

        prompt = prompts["article_prompt"].format(
            topic, "\n".join(str(sec) for sec in sections), related_links_prompt
        )

    article = generate_content(api_key, prompt, sections, model, temperature, presence_penalty, frequency_penalty, max_tokens)
    return article

class MyHTMLParser(HTMLParser):
    def __init__(self):
        super().__init__()
        self.text = []
        self.current_tag = None
        self.parent_tag = None
        self.current_href = None

    def handle_starttag(self, tag, attrs):
        self.current_tag = tag
        if tag in ["ul", "ol"]:
            self.parent_tag = tag
        if tag == "a":
            attrs_dict = dict(attrs)
            self.current_href = attrs_dict.get("href")

    def handle_endtag(self, tag):
        if tag == self.parent_tag:
            self.parent_tag = None
        if tag == "a":
            self.current_href = None
        self.current_tag = None

    def handle_data(self, data):
        if self.current_tag in ["h2", "h3", "h4", "p"]:
            self.text.append({"type": self.current_tag, "content": data.strip()})
        elif self.current_tag == "li":
            self.text.append({"type": self.current_tag, "content": data.strip(), "parent": self.parent_tag})
        elif self.current_tag == "a":
            self.text.append({"type": self.current_tag, "content": data.strip(), "href": self.current_href, "parent": self.parent_tag})

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    r = paragraph.add_run()
    r._r.append(hyperlink)

    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
    r.font.underline = True

    return hyperlink

def save_article_as_docx(filename, title, definition, content):
    parser = MyHTMLParser()
    parser.feed(content)
    parsed_content = parser.text

    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(definition)

    for item in parsed_content:
        if item["type"] in ["h2", "h3", "h4"]:
            level = int(item["type"][1])
            doc.add_heading(item["content"], level=level)
        elif item["type"] == "p":
            p = doc.add_paragraph(item["content"])
        elif item["type"] == "li":
            style = "ListBullet" if item["parent"] == "ul" else "ListNumber"
            p = doc.add_paragraph(item["content"], style=style)
        elif item["type"] == "a":
            add_hyperlink(p, item["content"], item["href"])

    doc.save(filename)
